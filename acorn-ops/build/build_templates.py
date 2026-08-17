"""
Generates the Word merge templates in src/word/.

Each is an ordinary .docx containing {{Token}} placeholders that modDocuments
substitutes at generation time. They are meant to be edited - open one in Word,
change the wording or drop the company logo in the header, and keep the tokens.

    python3 build_templates.py [output_dir]
"""

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

FONT = "Arial"
INK = RGBColor(0x1B, 0x2A, 0x22)
MUTED = RGBColor(0x5A, 0x6B, 0x62)
ACCENT = RGBColor(0x2F, 0x4A, 0x3C)


def base_document() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(9.5)
    style.font.color.rgb = INK
    style.paragraph_format.space_after = Pt(4)

    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
    return doc


def shade(cell, hex_colour):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_colour)
    cell._tc.get_or_add_tcPr().append(el)


def run(paragraph, text, *, size=9.5, bold=False, italic=False, colour=INK):
    r = paragraph.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = colour
    return r


def heading(doc, text, size=16):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run(p, text, size=size, bold=True, colour=ACCENT)
    return p


def sub(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run(p, text, size=8.5, italic=True, colour=MUTED)
    return p


def section_bar(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    shade(cell, "2F4A3C")
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    run(cell.paragraphs[0], text, size=10, bold=True, colour=RGBColor(0xFF, 0xFF, 0xFF))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def field_table(doc, rows, label_width=Cm(5.2)):
    """Two-column label / value grid - the shape most of these documents are."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, value) in enumerate(rows):
        left, right = table.rows[i].cells
        left.width = label_width
        shade(left, "F2F5F3")
        left.paragraphs[0].paragraph_format.space_after = Pt(0)
        right.paragraphs[0].paragraph_format.space_after = Pt(0)
        run(left.paragraphs[0], label, size=9, bold=True)
        run(right.paragraphs[0], value, size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def signature_block(doc, parties):
    table = doc.add_table(rows=2, cols=len(parties))
    table.style = "Table Grid"
    for i, party in enumerate(parties):
        head = table.rows[0].cells[i]
        shade(head, "F2F5F3")
        head.paragraphs[0].paragraph_format.space_after = Pt(0)
        run(head.paragraphs[0], party["title"], size=9, bold=True)
        body = table.rows[1].cells[i]
        for line in party["lines"]:
            p = body.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            run(p, line, size=9)
        body.paragraphs[0]._p.getparent().remove(body.paragraphs[0]._p)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def company_header(doc, right_text=""):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run(p, "{{Config.CompanyName}}", size=12, bold=True, colour=ACCENT)
    p2 = header.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    run(p2, "{{Config.CompanyAddress}}  {{Config.CompanyPostcode}}   "
            "T {{Config.CompanyPhone}}   E {{Config.CompanyEmail}}", size=8, colour=MUTED)
    p3 = header.add_paragraph()
    run(p3, "Waste carrier registration {{Config.CarrierLicence}}   "
            "Environmental permit {{Config.EnvironmentalPermit}}   "
            "VAT {{Config.CompanyVATNumber}}", size=8, colour=MUTED)
    if right_text:
        p3.add_run("   " + right_text)


def footer(doc, note):
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run(p, note, size=7.5, colour=MUTED)


def notes(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run(p, line, size=8, colour=MUTED)


# ---------------------------------------------------------------- templates

def waste_transfer_note(hazardous: bool):
    doc = base_document()
    company_header(doc)
    title = ("Hazardous Waste Consignment Note" if hazardous else "Waste Transfer Note")
    heading(doc, title)
    sub(doc, "Reference {{WTN.WTNRef}}    Issued {{WTN.IssueDate}}    Job {{Job.JobID}}")

    section_bar(doc, "A   Description of the waste")
    field_table(doc, [
        ("EWC / List of Waste code", "{{WTN.EWCCode}}"),
        ("Description of the waste", "{{WTN.WasteDescription}}"),
        ("How it is contained", "{{WTN.ContainerType}}"),
        ("Quantity (tonnes)", "{{WTN.QuantityTonnes}}"),
        ("Producer's SIC code (2007)", "{{WTN.SICCode}}"),
    ])

    section_bar(doc, "B   The person transferring the waste (the producer / current holder)")
    field_table(doc, [
        ("Name", "{{Customer.CustomerName}}"),
        ("Address", "{{Customer.Address1}} {{Customer.Address2}}, {{Customer.Town}} "
                    "{{Customer.Postcode}}"),
        ("Site the waste came from", "{{Site.SiteName}}, {{Site.Address1}}, {{Site.Town}} "
                                     "{{Site.Postcode}}"),
        ("Account reference", "{{Customer.CustomerID}}"),
    ])

    section_bar(doc, "C   The person receiving the waste (the carrier and the site)")
    field_table(doc, [
        ("Name", "{{Config.CompanyName}}"),
        ("Address", "{{Config.CompanyAddress}}, {{Config.CompanyPostcode}}"),
        ("Waste carrier registration", "{{Config.CarrierLicence}}"),
        ("Environmental permit", "{{Config.EnvironmentalPermit}}"),
        ("Capacity", "Carrier and operator of the receiving facility"),
    ])

    if hazardous:
        section_bar(doc, "D   Consignor's certificate")
        p = doc.add_paragraph()
        run(p, "I certify that I have fulfilled my duty to apply the waste hierarchy as required "
               "by regulation 12 of the Waste (England and Wales) Regulations 2011. The "
               "information given below is correct, the carrier is registered, and the waste is "
               "packaged and labelled correctly for carriage.", size=8.5)
        doc.add_paragraph()

    section_bar(doc, ("E   Signatures" if hazardous else "D   Signatures"))
    signature_block(doc, [
        {"title": "Transferor (producer)",
         "lines": ["Name: {{WTN.ProducerSignedBy}}", "Signature: ..............................",
                   "Date: {{WTN.ProducerSignedOn}}", "On behalf of: {{Customer.CustomerName}}"]},
        {"title": "Transferee (carrier)",
         "lines": ["Name: {{WTN.CarrierSignedBy}}", "Signature: ..............................",
                   "Date: {{WTN.IssueDate}}", "On behalf of: {{Config.CompanyName}}"]},
    ])

    p = doc.add_paragraph()
    run(p, "Place of transfer: {{Site.SiteName}}, {{Site.Postcode}}      "
           "Time of transfer: ..................", size=9)

    doc.add_paragraph()
    retention = ("three years" if hazardous else "two years")
    notes(doc, [
        f"Both parties must keep a copy of this note for {retention} and produce it on request "
        f"by the regulator.",
        "Section 34 of the Environmental Protection Act 1990 places a duty of care on everyone "
        "who produces, carries, keeps, treats or disposes of controlled waste.",
        "Generated {{Now}} by {{User}}. Retention expires {{WTN.RetentionUntil}}.",
    ])
    footer(doc, f"{title} {{{{WTN.WTNRef}}}} - {{{{Config.CompanyName}}}} - page ")
    return doc


def job_confirmation():
    doc = base_document()
    company_header(doc)
    heading(doc, "Job Confirmation")
    sub(doc, "{{Job.JobID}}    Confirmed {{Today}}")

    p = doc.add_paragraph()
    run(p, "{{Customer.ContactName}}", size=10)
    p = doc.add_paragraph()
    run(p, "Thank you for your order. We have booked the following:", size=9.5)
    doc.add_paragraph()

    field_table(doc, [
        ("Our reference", "{{Job.JobID}}"),
        ("Your order number", "{{Job.PONumber}}"),
        ("Service", "{{Job.ServiceType}}"),
        ("Container", "{{Job.ContainerType}}"),
        ("Delivery site", "{{Site.SiteName}}, {{Site.Address1}}, {{Site.Town}} {{Site.Postcode}}"),
        ("Access notes", "{{Site.AccessNotes}}"),
        ("Wanted for", "{{Job.RequestedDate}}"),
        ("Scheduled for", "{{Job.ScheduledDate}}"),
        ("Price agreed (excl. VAT)", "£{{Job.PriceAgreed}}"),
        ("Waste type", "{{Job.WasteDescription}} (EWC {{Job.EWCCode}})"),
    ])

    section_bar(doc, "What we cannot take")
    notes(doc, [
        "Plasterboard, asbestos, tyres, fridges, mattresses, gas cylinders, batteries, paint, "
        "oils, solvents, clinical waste and any liquid waste must not go in the container. They "
        "are chargeable separately and a contaminated load may be rejected at the facility.",
        "Please do not load above the rim. An overloaded container cannot legally be moved and "
        "will be left on site.",
        "Where the container stands on the public highway a permit is required. Permit reference "
        "for this job: {{Job.PermitRef}}",
    ])

    section_bar(doc, "Duty of care")
    notes(doc, [
        "A waste transfer note will be raised for this movement and a copy sent to you. "
        "{{Config.CompanyName}} is a registered waste carrier, registration "
        "{{Config.CarrierLicence}}, and operates under environmental permit "
        "{{Config.EnvironmentalPermit}}.",
    ])

    doc.add_paragraph()
    p = doc.add_paragraph()
    run(p, "{{User}}", size=9.5, bold=True)
    p = doc.add_paragraph()
    run(p, "{{Config.CompanyName}}   {{Config.CompanyPhone}}   {{Config.CompanyEmail}}", size=9)

    footer(doc, "Job confirmation {{Job.JobID}} - generated {{Now}}")
    return doc


def quotation():
    doc = base_document()
    company_header(doc)
    heading(doc, "Quotation")
    sub(doc, "{{Job.JobID}}    {{Today}}    Valid for 30 days from the date above")

    field_table(doc, [
        ("For the attention of", "{{Customer.ContactName}}"),
        ("Company", "{{Customer.CustomerName}}"),
        ("Site", "{{Site.SiteName}}, {{Site.Town}} {{Site.Postcode}}"),
        ("Our reference", "{{Job.JobID}}"),
    ])

    section_bar(doc, "What we are quoting for")
    field_table(doc, [
        ("Service", "{{Job.ServiceType}}"),
        ("Container", "{{Job.ContainerType}}"),
        ("Waste type", "{{Job.WasteDescription}} (EWC {{Job.EWCCode}})"),
        ("Price, excluding VAT", "£{{Job.PriceAgreed}}"),
    ])

    section_bar(doc, "What the price includes, and what it does not")
    notes(doc, [
        "The price covers delivery, the hire period stated on our current rate card, collection, "
        "and the tonnage allowance for the container size quoted.",
        "Tonnage over the allowance is charged at the rate on the rate card, worked out from our "
        "weighbridge ticket. A copy of the ticket is available on request.",
        "Hire beyond the included period is charged per day.",
        "Prohibited items - plasterboard, asbestos, tyres, fridges, mattresses, gas cylinders, "
        "batteries, paint, oils, solvents, clinical and liquid waste - are chargeable separately.",
        "Where a highway permit is needed, the permit fee is charged at cost.",
        "Prices exclude VAT, which is charged at the prevailing rate.",
    ])

    section_bar(doc, "Our permissions")
    notes(doc, [
        "Registered waste carrier {{Config.CarrierLicence}}. Environmental permit "
        "{{Config.EnvironmentalPermit}}. We will issue a waste transfer note for every "
        "movement and can report the recovery rate achieved on your waste.",
    ])

    doc.add_paragraph()
    p = doc.add_paragraph()
    run(p, "Quoted by {{User}} on {{Today}}", size=9)
    footer(doc, "Quotation {{Job.JobID}} - {{Config.CompanyName}}")
    return doc


def duty_of_care_statement():
    doc = base_document()
    company_header(doc)
    heading(doc, "Duty of Care Statement")
    sub(doc, "{{Customer.CustomerName}}    Account {{Customer.CustomerID}}    Issued {{Today}}")

    p = doc.add_paragraph()
    run(p, "This statement confirms how the waste collected from you has been handled, and where "
           "it went. It is provided so that you can evidence your own duty of care under section "
           "34 of the Environmental Protection Act 1990.", size=9.5)
    doc.add_paragraph()

    section_bar(doc, "Our permissions")
    field_table(doc, [
        ("Company", "{{Config.CompanyName}}"),
        ("Waste carrier registration", "{{Config.CarrierLicence}}"),
        ("Environmental permit", "{{Config.EnvironmentalPermit}}"),
        ("Site address", "{{Config.CompanyAddress}}, {{Config.CompanyPostcode}}"),
    ])

    section_bar(doc, "Period covered and tonnage")
    notes(doc, [
        "Complete the table below from the Dashboard before issuing. The figures come from "
        "weighbridge tickets, not estimates.",
    ])
    field_table(doc, [
        ("Period", ".............................................."),
        ("Total tonnage collected", ".............................................."),
        ("Tonnage recycled or recovered", ".............................................."),
        ("Tonnage to landfill", ".............................................."),
        ("Recovery rate achieved", ".............................................."),
    ])

    section_bar(doc, "How the waste was treated")
    notes(doc, [
        "Material is received at our permitted transfer station, segregated and sent to "
        "authorised outlets. Every outlet's permit is verified before we use it and reviewed "
        "annually; records are held on our approved outlet register.",
        "The waste hierarchy is applied: material is prepared for reuse where possible, then "
        "recycled, then recovered, with landfill used only for the residue that cannot be "
        "treated any other way.",
        "A waste transfer note was raised for every movement and is retained for two years "
        "(three years for hazardous consignments).",
    ])

    doc.add_paragraph()
    p = doc.add_paragraph()
    run(p, "Signed: ..............................   Name: {{User}}   Date: {{Today}}", size=9)
    footer(doc, "Duty of care statement - {{Customer.CustomerID}} - generated {{Now}}")
    return doc


def ncr_report():
    doc = base_document()
    company_header(doc)
    heading(doc, "Nonconformity and Corrective Action Report")
    sub(doc, "{{NCR.NCRID}}    Raised {{NCR.RaisedOn}} by {{NCR.RaisedBy}}")

    section_bar(doc, "1   What happened")
    field_table(doc, [
        ("Source", "{{NCR.Source}}"),
        ("Type", "{{NCR.Type}}"),
        ("Severity", "{{NCR.Severity}}"),
        ("Linked job", "{{NCR.LinkedJobID}}"),
        ("Linked customer", "{{NCR.LinkedCustomerID}}"),
        ("Description", "{{NCR.Description}}"),
    ])

    section_bar(doc, "2   Immediate correction")
    field_table(doc, [("What was done straight away", "{{NCR.ImmediateCorrection}}")])

    section_bar(doc, "3   Root cause")
    notes(doc, [
        "State why it happened, not what happened. If the answer is 'human error', keep asking "
        "why until you reach something that can actually be changed.",
    ])
    field_table(doc, [("Root cause", "{{NCR.RootCause}}")])

    section_bar(doc, "4   Corrective action")
    field_table(doc, [
        ("Action to stop it recurring", "{{NCR.CorrectiveAction}}"),
        ("Owner", "{{NCR.ActionOwner}}"),
        ("Due", "{{NCR.DueDate}}"),
    ])

    section_bar(doc, "5   Verification of effectiveness")
    notes(doc, [
        "An action is not closed when it is done - it is closed when someone has checked that it "
        "worked. Say how it was checked, by whom, and when.",
    ])
    field_table(doc, [
        ("Evidence it worked", "{{NCR.EffectivenessCheck}}"),
        ("Closed on", "{{NCR.ClosedOn}}"),
        ("Status", "{{NCR.Status}}"),
        ("Days open", "{{NCR.DaysOpen}}"),
    ])

    doc.add_paragraph()
    p = doc.add_paragraph()
    run(p, "Verified by: ..............................   Date: ..................", size=9)
    footer(doc, "NCR {{NCR.NCRID}} - ISO 9001 / 14001 / 45001 clause 10.2 - generated {{Now}}")
    return doc


def toolbox_talk():
    doc = base_document()
    company_header(doc)
    heading(doc, "Toolbox Talk Record")
    sub(doc, "Date {{Date}}    Delivered by {{User}}")

    field_table(doc, [
        ("Subject", ".............................................................."),
        ("Location", ".............................................................."),
        ("Duration", ".............................................................."),
        ("Reason for the talk", "..............................................................."),
    ])

    section_bar(doc, "Key points covered")
    for _ in range(6):
        p = doc.add_paragraph()
        run(p, "•  ..........................................................................."
               "..............................", size=9)

    section_bar(doc, "Attendance")
    table = doc.add_table(rows=13, cols=3)
    table.style = "Table Grid"
    for i, head in enumerate(("Name", "Signature", "Job role")):
        cell = table.rows[0].cells[i]
        shade(cell, "F2F5F3")
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        run(cell.paragraphs[0], head, size=9, bold=True)

    doc.add_paragraph()
    notes(doc, [
        "File the signed original in 05_Compliance\\Training and add a row to the Training sheet "
        "of the Compliance workbook so it shows on the training matrix.",
        "ISO 45001 clause 7.3 - workers must be aware of the hazards and of their own "
        "contribution to the effectiveness of the management system.",
    ])
    footer(doc, "Toolbox talk record - {{Config.CompanyName}} - generated {{Now}}")
    return doc


def site_induction():
    doc = base_document()
    company_header(doc)
    heading(doc, "Site Induction Record")
    sub(doc, "{{Name}}    {{Date}}")

    field_table(doc, [
        ("Name", "{{Name}}"),
        ("Company", "................................................."),
        ("Role / reason on site", "................................................."),
        ("Vehicle registration", "................................................."),
        ("Inducted by", "{{User}}"),
        ("Date", "{{Date}}"),
    ])

    section_bar(doc, "Covered during the induction")
    items = [
        "Site layout, one-way system and speed limit",
        "Pedestrian routes and the segregation from vehicle movements",
        "Where to report on arrival and on departure",
        "Mandatory PPE: hi-vis, safety boots, hard hat, gloves, eye protection",
        "Exclusion zones around tipping, loading and the picking line",
        "Weighbridge procedure - stay in the cab unless told otherwise",
        "What must not be tipped, and what to do if a prohibited item is found",
        "Fire points, muster point and the alarm signal",
        "First aid, accident and near miss reporting",
        "Spill kit locations and what to do about a spill",
        "Permitted waste types and the site permit conditions",
        "No smoking, no mobile phones outside the cab, no lone working in the pit",
    ]
    for item in items:
        p = doc.add_paragraph()
        run(p, f"☐   {item}", size=9)

    doc.add_paragraph()
    signature_block(doc, [
        {"title": "Person inducted",
         "lines": ["I confirm I have received and understood this induction.",
                   "Signature: ..............................", "Date: {{Date}}"]},
        {"title": "Inducted by",
         "lines": ["Name: {{User}}", "Signature: ..............................",
                   "Date: {{Date}}"]},
    ])
    footer(doc, "Site induction record - {{Config.CompanyName}} - generated {{Now}}")
    return doc


TEMPLATES = {
    "WasteTransferNote.docx": lambda: waste_transfer_note(False),
    "HazardousConsignmentNote.docx": lambda: waste_transfer_note(True),
    "JobConfirmation.docx": job_confirmation,
    "Quotation.docx": quotation,
    "DutyOfCareStatement.docx": duty_of_care_statement,
    "NCRReport.docx": ncr_report,
    "ToolboxTalkRecord.docx": toolbox_talk,
    "SiteInductionRecord.docx": site_induction,
}


def main():
    outdir = Path(sys.argv[1] if len(sys.argv) > 1 else Path(__file__).parent.parent / "src" / "word")
    outdir.mkdir(parents=True, exist_ok=True)
    for name, builder in TEMPLATES.items():
        doc = builder()
        doc.save(outdir / name)
        print(f"  wrote {name}")


if __name__ == "__main__":
    main()
